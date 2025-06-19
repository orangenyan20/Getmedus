import streamlit as st
import requests
from bs4 import BeautifulSoup
import re
import time
from io import BytesIO
from docx import Document
from docx.shared import Inches

# ファイル読み込み（BOM除去 + 多エンコーディング対応）
def try_read_file(uploaded_file):
    raw_bytes = uploaded_file.read()
    if raw_bytes.startswith(b'\xff\xfe'):
        encoding = 'utf-16'
    elif raw_bytes.startswith(b'\xfe\xff'):
        encoding = 'utf-16'
    elif raw_bytes.startswith(b'\xef\xbb\xbf'):
        encoding = 'utf-8-sig'
    else:
        encodings = ['utf-8', 'shift_jis', 'cp932', 'iso-2022-jp', 'utf-16', 'utf-16-le', 'utf-16-be']
        for enc in encodings:
            try:
                return [line.strip() for line in raw_bytes.decode(enc).splitlines() if line.strip()]
            except:
                continue
        st.error("❌ ファイルの読み込みに失敗しました。文字コードを確認してください。")
        return []
    try:
        return [line.strip() for line in raw_bytes.decode(encoding).splitlines() if line.strip()]
    except:
        st.error("❌ BOM付きファイルの読み込みに失敗しました。")
        return []

# 問題番号からURL生成
def generate_urls_from_ids(question_ids):
    base_url = "https://medu4.com/"
    return [f"{base_url}{qid.strip()}" for qid in question_ids if qid.strip()]

# ページデータ取得（正答率追加）
def get_page_text(url, get_images=True):
    try:
        response = requests.get(url)
        if response.status_code != 200:
            return None

        soup = BeautifulSoup(response.text, 'html.parser')

        category = soup.find('span', class_='button-small-line')
        category_name = category.text.strip() if category else '分野名なし'

        problem = soup.find('div', class_='quiz-body mb-64')
        problem_text = problem.text.strip() if problem else '問題文なし'

        choices = []
        for choice in soup.find_all('div', class_='box-select'):
            choice_header = choice.find('span', {'class': 'choice-header'}).text.strip()
            choice_text = choice.find_all('span')[1].text.strip()
            choices.append(f"{choice_header} {choice_text}")

        h4_tags = soup.find_all('h4')
        answer_text = '解答なし'
        question_id = '問題番号なし'
        if len(h4_tags) >= 2:
            answer_text = h4_tags[0].text.strip()
            question_id_match = re.search(r'([0-9]{3}[A-Za-z][0-9]+)', h4_tags[1].text)
            if question_id_match:
                question_id = question_id_match.group(1)

        explanation = soup.find('div', class_='explanation')
        explanation_text = explanation.text.strip() if explanation else '解説なし'

        image_urls = []
        if get_images:
            image_divs = soup.find_all('div', class_='box-quiz-image mb-32')
            for div in image_divs:
                for img_tag in div.find_all('img'):
                    if img_tag.get('src'):
                        img_url = img_tag['src'].replace('thumb_', '')
                        image_urls.append(img_url)

        # 正答率
        accuracy = '正答率不明'
        for p in soup.find_all('p', class_='commentary-date'):
            match = re.search(r'正答率：(\d+)%', p.text)
            if match:
                accuracy = int(match.group(1))
                break

        return {
            "category": category_name,
            "problem": problem_text,
            "choices": choices,
            "answer": answer_text,
            "question_id": question_id,
            "explanation": explanation_text,
            "images": image_urls,
            "accuracy": accuracy
        }
    except Exception as e:
        return None

# Wordファイル作成
def create_word_doc(pages_data, search_query, include_images=True):
    doc = Document()
    doc.add_heading('検索結果', 0)
    doc.add_paragraph(f"取得問題数: {len(pages_data)}問")

    for idx, page_data in enumerate(pages_data, start=1):
        title = f"問題{idx} {page_data['question_id']}"
        doc.add_paragraph(title, style='Heading2')
        doc.add_paragraph(f"正答率: {page_data['accuracy']}%")
        doc.add_paragraph(f"問題文: {page_data['problem']}")

        if include_images and page_data['images']:
            for img_url in page_data['images']:
                try:
                    response = requests.get(img_url)
                    if response.status_code == 200:
                        image_stream = BytesIO(response.content)
                        doc.add_paragraph()
                        doc.add_picture(image_stream, width=Inches(2.5))
                    else:
                        doc.add_paragraph(f"画像取得失敗: {img_url}")
                except Exception as e:
                    doc.add_paragraph(f"画像取得中エラー: {e}")

        doc.add_paragraph("選択肢:")
        for choice in page_data['choices']:
            doc.add_paragraph(choice)
        doc.add_paragraph(f"{page_data['answer']}")
        doc.add_paragraph(f"解説: {page_data['explanation']}")
        doc.add_paragraph("-" * 50)

    filename = f"{search_query}_search_results.docx"
    doc.save(filename)
    return filename

# Streamlit UI
st.title("🩺 Medu4 問題番号から収集ツール")

uploaded_file = st.file_uploader("📄 問題番号のファイルをアップロード（.txt / .csv）", type=["txt", "csv"])
include_images = st.checkbox("🖼️ 画像も含める", value=True)
sort_by_accuracy = st.checkbox("📉 正答率が低い順に並び替える", value=False)

if uploaded_file:
    question_ids = try_read_file(uploaded_file)
    if not question_ids:
        st.stop()

    urls = generate_urls_from_ids(question_ids)

    st.write(f"{len(urls)}個の問題を取得します。")
    progress_bar = st.progress(0)
    pages_data = []

    for i, url in enumerate(urls):
        page_data = get_page_text(url, get_images=include_images)
        if page_data:
            pages_data.append(page_data)
        else:
            st.warning(f"❌ URL取得失敗: {url}")
        progress_bar.progress((i + 1) / len(urls))
        time.sleep(0.2)

    # 並び替え
    if sort_by_accuracy:
        pages_data.sort(key=lambda x: (x["accuracy"] if isinstance(x["accuracy"], int) else 9999))

    with st.spinner("Wordファイルを作成中..."):
        filename = create_word_doc(pages_data, "問題番号リスト", include_images=include_images)

    st.success("✅ Wordファイルの生成が完了しました！")
    with open(filename, "rb") as file:
        st.download_button("📄 Wordファイルをダウンロード", file, file_name=filename)
